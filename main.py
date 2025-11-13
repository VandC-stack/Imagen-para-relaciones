import os
import json
import sys
import re
import unicodedata
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Inches
from PIL import Image

# ------------------------------------------------------------
# REGISTRO DE FALLOS
# ------------------------------------------------------------
try:
    from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro
except Exception as e:
    print(f"Error al importar registro_fallos: {e}")
    registrar_fallo = None
    limpiar_registro = None
    mostrar_registro = None

# ------------------------------------------------------------
# CONFIGURACIÓN
# ------------------------------------------------------------
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

CONFIG_FILE = os.path.abspath("config.json")

FORBIDDEN_TOKENS = {
    "TOTAL", "CANTIDAD", "FACTURA", "MARCA", "DESCRIPCION", "DESCRIPCIÓN",
    "FECHA", "CONTRATO", "PRESENTACION", "PRESENTACIÓN", "SISTEMA",
    "ALEATORIO", "DICTAMEN", "PRODUCTO", "RELACION", "RELACIÓN",
    "MODELO", "ORIGEN", "CHINA", "MALASIA", "ALEMANIA", "RUMANIA",
    "ITALIA", "BRASIL", "DIMENSIONES", "CONTENIDO", "ETIQUETA", "CONTEN",
}
IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

# ------------------------------------------------------------
# CONFIG
# ------------------------------------------------------------
def guardar_config(data):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def cargar_config():
    if not os.path.exists(CONFIG_FILE):
        return {}
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def seleccionar_carpeta(titulo):
    Tk().withdraw()
    carpeta = filedialog.askdirectory(title=titulo)
    return carpeta.replace("\\", "/").strip() if carpeta else None

def obtener_rutas():
    cfg = cargar_config()
    ruta_docs = cfg.get("ruta_docs", "")
    ruta_imgs = cfg.get("ruta_imagenes", "")

    if not os.path.isdir(ruta_docs):
        ruta_docs = seleccionar_carpeta("Selecciona la carpeta de documentos .docx")
        if not ruta_docs:
            return None, None
        cfg["ruta_docs"] = ruta_docs

    if not os.path.isdir(ruta_imgs):
        ruta_imgs = seleccionar_carpeta("Selecciona la carpeta de imágenes")
        if not ruta_imgs:
            return None, None
        cfg["ruta_imagenes"] = ruta_imgs

    guardar_config(cfg)
    return ruta_docs, ruta_imgs

# ------------------------------------------------------------
# UTILIDADES
# ------------------------------------------------------------
def _sin_acentos(s):
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")

def normalizar_cadena_alnum_mayus(s):
    return re.sub(r"[^A-Za-z0-9]", "", s or "").upper()

def contiene_digito(s):
    return any(c.isdigit() for c in s or "")

# ------------------------------------------------------------
# INDEXADO DE IMÁGENES (modo simple)
# ------------------------------------------------------------
def indexar_imagenes(carpeta_imagenes):
    index = []
    for nombre in os.listdir(carpeta_imagenes):
        base, ext = os.path.splitext(nombre)
        if ext.lower() not in IMG_EXTS:
            continue
        index.append({
            "name": nombre,
            "base": base,
            "ext": ext,
            "base_norm": normalizar_cadena_alnum_mayus(base),
            "path": os.path.join(carpeta_imagenes, nombre),
        })
    return index

def norm_path_key(path):
    return os.path.normcase(os.path.normpath(path or ""))

def buscar_imagen_index(index, codigo_canonico, usadas_paths, usadas_bases):
    code = normalizar_cadena_alnum_mayus(codigo_canonico)
    if not code:
        return None

    exactos = [
        it for it in index
        if it["base_norm"] == code
        and norm_path_key(it["path"]) not in usadas_paths
        and it["base_norm"] not in usadas_bases
    ]
    if exactos:
        return exactos[0]["path"]

    parciales = [
        it for it in index
        if (code in it["base_norm"] or it["base_norm"] in code)
        and norm_path_key(it["path"]) not in usadas_paths
        and it["base_norm"] not in usadas_bases
    ]
    if not parciales:
        return None

    def score(it):
        bn = it["base_norm"]
        starts = bn.startswith(code)
        ends = bn.endswith(code)
        delta = abs(len(bn) - len(code))
        return (0 if starts or ends else 1, delta, bn)

    parciales.sort(key=score)
    return parciales[0]["path"]

# ------------------------------------------------------------
# EXTRACCIÓN DE CÓDIGOS
# ------------------------------------------------------------
def extraer_codigos(doc):
    codigos = []
    patron_general = re.compile(r"[A-Za-z0-9][A-Za-z0-9.\-]{4,}", re.IGNORECASE)
    patron_bosch = re.compile(r"(?:No\.?\s*)?(\d(?:\s?\d){8,12})")

    for tabla in doc.tables:
        if not tabla.rows:
            continue

        idx_codigo = None
        for r in range(min(3, len(tabla.rows))):
            for i, celda in enumerate(tabla.rows[r].cells):
                t_norm = _sin_acentos(celda.text or "").upper()
                if ("CODIGO" in t_norm or "SKU" in t_norm or "CLAVE" in t_norm):
                    idx_codigo = i
                    break
            if idx_codigo is not None:
                break

        columnas = [idx_codigo] if idx_codigo is not None else range(len(tabla.rows[0].cells))

        for fila in tabla.rows[1:]:
            for j in columnas:
                texto = (fila.cells[j].text or "").strip()

                for m in patron_bosch.findall(texto):
                    num = re.sub(r"\s+", "", m)
                    if 8 <= len(num) <= 13:
                        codigos.append(num)

                for m in patron_general.findall(texto):
                    canon = normalizar_cadena_alnum_mayus(m)
                    if canon and contiene_digito(canon) and canon not in FORBIDDEN_TOKENS:
                        codigos.append(canon)

    return list(dict.fromkeys(codigos))

# ------------------------------------------------------------
# INSERCIÓN DE IMÁGENES (ambos modos)
# ------------------------------------------------------------
H_MAX_W_CM = 4.36
H_MAX_H_CM = 6.37
V_MAX_W_CM = 8.13
V_MAX_H_CM = 4.84

def insertar_imagen_con_transparencia(run, img_path):
    try:
        with Image.open(img_path) as img:
            w_px, h_px = img.size

        w_in, h_in = w_px / 96.0, h_px / 96.0

        if w_px >= h_px:
            max_w_in = H_MAX_W_CM / 2.54
            max_h_in = H_MAX_H_CM / 2.54
        else:
            max_w_in = V_MAX_W_CM / 2.54
            max_h_in = V_MAX_H_CM / 2.54

        scale = min(max_w_in / w_in, max_h_in / h_in, 1)
        new_w_in = w_in * scale
        new_h_in = h_in * scale

        if new_w_in <= max_w_in:
            run.add_picture(img_path, width=Inches(new_w_in))
        else:
            run.add_picture(img_path, height=Inches(new_h_in))

        run.add_text(" ")
    except Exception as e:
        print(f"Error al insertar {img_path}: {e}")

# ------------------------------------------------------------
# MODO NUEVO: CARPETAS POR CÓDIGO
# ------------------------------------------------------------
def insertar_imagenes_en_docx_carpetas(ruta_doc, carpeta_base):
    print(f"Procesando documento (modo carpetas): {ruta_doc}")
    doc = Document(ruta_doc)
    codigos = extraer_codigos(doc)

    if not codigos:
        if registrar_fallo:
            registrar_fallo(os.path.basename(ruta_doc))
        return

    imagen_insertada = False

    for p in doc.paragraphs:
        if "${etiqueta1}" in (p.text or ""):
            p.clear()
            run = p.add_run()

            for codigo in codigos:
                carpeta_codigo = os.path.join(carpeta_base, codigo)
                if not os.path.isdir(carpeta_codigo):
                    print(f"No existe carpeta para {codigo}")
                    continue

                for archivo in os.listdir(carpeta_codigo):
                    ext = os.path.splitext(archivo)[1].lower()
                    if ext in IMG_EXTS:
                        img_path = os.path.join(carpeta_codigo, archivo)
                        insertar_imagen_con_transparencia(run, img_path)
                        imagen_insertada = True
                        print(f"Imagen insertada: {img_path}")

            break

    if not imagen_insertada and registrar_fallo:
        registrar_fallo(os.path.basename(ruta_doc))

    doc.save(ruta_doc)
    print(f"Documento actualizado: {ruta_doc}")

# ------------------------------------------------------------
# MODO NORMAL (ACTUAL)
# ------------------------------------------------------------
def insertar_imagenes_en_docx(ruta_doc, carpeta_imagenes, index):
    print(f"Procesando documento: {ruta_doc}")
    doc = Document(ruta_doc)
    codigos = extraer_codigos(doc)

    if not codigos:
        if registrar_fallo:
            registrar_fallo(os.path.basename(ruta_doc))
        return

    usadas_paths, usadas_bases = set(), set()
    imagen_insertada = False

    for p in doc.paragraphs:
        if "${etiqueta1}" in (p.text or ""):
            p.clear()
            run = p.add_run()

            for codigo in codigos:
                img_path = buscar_imagen_index(index, codigo, usadas_paths, usadas_bases)
                if img_path:
                    usadas_paths.add(norm_path_key(img_path))
                    usadas_bases.add(normalizar_cadena_alnum_mayus(os.path.splitext(os.path.basename(img_path))[0]))
                    insertar_imagen_con_transparencia(run, img_path)
                    imagen_insertada = True
                else:
                    print(f"No se encontró imagen para {codigo}")

            break

    if not imagen_insertada and registrar_fallo:
        registrar_fallo(os.path.basename(ruta_doc))

    doc.save(ruta_doc)
    print(f"Documento actualizado: {ruta_doc}")

# ------------------------------------------------------------
# PROCESAMIENTO GENERAL
# ------------------------------------------------------------
def procesar_lote():
    if limpiar_registro:
        limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    archivos = [f for f in os.listdir(ruta_docs) if f.endswith(".docx") and not f.startswith("~$")]

    if not archivos:
        print(f"No se encontraron archivos .docx en '{ruta_docs}'")
        return

    cfg = cargar_config()
    modo = cfg.get("modo_pegado", "simple")

    if modo == "simple":
        index = indexar_imagenes(ruta_imgs)

    for archivo in archivos:
        ruta_doc = os.path.join(ruta_docs, archivo)

        if modo == "carpetas":
            insertar_imagenes_en_docx_carpetas(ruta_doc, ruta_imgs)
        else:
            insertar_imagenes_en_docx(ruta_doc, ruta_imgs, index)

    if mostrar_registro:
        mostrar_registro()

    log = os.path.abspath("documentos_sin_imagenes.txt")
    if os.path.exists(log):
        os.startfile(log)

# ------------------------------------------------------------
if __name__ == "__main__":
    procesar_lote()