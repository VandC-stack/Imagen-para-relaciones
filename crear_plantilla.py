from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
import os

def agregar_fondo_como_encabezado(doc, image_path):
    """Agrega una imagen de fondo a través del encabezado (detrás del texto)."""
    section = doc.sections[0]
    header = section.header
    
    # Limpiar el header primero
    for element in header._element:
        header._element.remove(element)
    
    paragraph = header.add_paragraph()

    if not os.path.exists(image_path):
        print(f"⚠️ No se encontró la imagen: {image_path}")
        return

    # Tamaño total de la página
    ancho_pagina = section.page_width
    alto_pagina = section.page_height

    # Crear relación para la imagen
    rel_id = header.part.relate_to(
        image_path,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        is_external=True
    )

    # Crear XML de imagen flotante
    background_xml = f'''
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
             xmlns:v="urn:schemas-microsoft-com:vml"
             xmlns:o="urn:schemas-microsoft-com:office:office"
             xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
        <v:shape id="Fondo" 
                 style="position:absolute;margin-left:0;margin-top:0;width:{ancho_pagina.pt}pt;height:{alto_pagina.pt}pt;z-index:-251659264" 
                 o:allowincell="f" 
                 o:spid="_x0000_s1025" 
                 type="#_x0000_t75">
            <v:imagedata r:id="{rel_id}" o:title="Fondo"/>
        </v:shape>
    </w:pict>
    '''

    try:
        background_element = parse_xml(background_xml)
        paragraph._p.append(background_element)
        print(f"✅ Fondo agregado correctamente desde: {image_path}")
    except Exception as e:
        print(f"❌ Error al agregar fondo: {e}")

def crear_plantilla_con_fondo():
    """Crea la plantilla de dictamen con el diseño especificado."""
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Imagen de fondo
    image_path = "img/Fondo.jpeg"
    
    if not os.path.exists(image_path):
        print(f"❌ No se puede encontrar la imagen de fondo: {image_path}")
        print("💡 Asegúrate de que el archivo 'Fondo.jpeg' esté en la carpeta 'img/'")
        return
    
    try:
        agregar_fondo_como_encabezado(doc, image_path)
    except Exception as e:
        print(f"❌ Error al agregar fondo: {e}")

    # ==================== TÍTULO ====================
    titulo = doc.add_paragraph()
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = titulo.add_run("DICTAMEN DE CUMPLIMIENTO")
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(17)
    run.font.bold = True

    subtitulo = doc.add_paragraph("Unidad de Inspección de Etiquetado NOM-024")
    subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitulo.runs[0]
    run.font.name = "Helvetica Neue"
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ==================== FECHAS CON CUADROS ====================
    # Crear tabla para las fechas con bordes
    table_fechas = doc.add_table(rows=2, cols=2)
    table_fechas.style = "Table Grid"
    
    # Configurar ancho de columnas
    for row in table_fechas.rows:
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(3)
    
    # Primera fila - Encabezados
    table_fechas.cell(0, 0).text = "Fecha de Inspección"
    table_fechas.cell(0, 1).text = "{{ fverificacion }}"
    
    # Segunda fila - Valores
    table_fechas.cell(1, 0).text = "Fecha de Emisión"
    table_fechas.cell(1, 1).text = "{{ femision }}"
    
    # Formatear todas las celdas de la tabla de fechas
    for row in table_fechas.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Helvetica Neue"
                    run.font.size = Pt(10)
                    if row == table_fechas.rows[0]:  # Solo la primera fila en negrita
                        run.font.bold = True

    doc.add_paragraph()

    # ==================== CLIENTE Y RFC ====================
    p = doc.add_paragraph()
    p.add_run("Cliente: ").bold = True
    p.add_run("{{ cliente }}")
    for run in p.runs:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("RFC: ").bold = True
    p.add_run("{{ rfc }}")
    for run in p.runs:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ==================== TEXTO PRINCIPAL ====================
    texto = (
        "De conformidad en lo dispuesto en los artículos 53, 56 fracción I, 60 fracción I, 62, 64, 68 y 140 de la Ley de Infraestructura de la "
        "Calidad; 50 del Reglamento de la Ley Federal de Metrología y Normalización; Punto 2.4.8 Fracción III ACUERDO por el que la "
        "Secretaría de Economía emite Reglas y criterios de carácter general en materia de comercio exterior; publicado en el Diario Oficial de la "
        "Federación el 09 de mayo de 2022 y posteriores modificaciones; esta Unidad de Inspección a solicitud de la persona moral denominada "
        "{{ cliente }} dictamina el Producto: {{ producto }}; que la mercancía importada bajo el pedimento aduanal No. {{ pedimento }} de fecha "
        "{{ fverificacionlarga }}, fue etiquetada conforme a los requisitos de Información Comercial en el capítulo {{ capitulo }} de la Norma Oficial Mexicana "
        "{{ norma }} {{ normades }} Cualquier otro requisito establecido en la norma referida, es responsabilidad del titular de este Dictamen."
    )
    p = doc.add_paragraph(texto)
    for run in p.runs:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ==================== TABLA DE PRODUCTOS ====================
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    
    headers = ["MARCA", "CÓDIGO", "FACTURA", "CANTIDAD"]
    values = ["{{ rowMarca }}", "{{ rowCodigo }}", "{{ rowFactura }}", "{{ rowCantidad }}"]

    # Encabezados de tabla
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Helvetica Neue"

    # Valores de tabla
    for i, value in enumerate(values):
        cell = table.rows[1].cells[i]
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.name = "Helvetica Neue"

    # ==================== TAMAÑO DEL LOTE ====================
    table_lote = doc.add_table(rows=1, cols=2)
    table_lote.style = "Table Grid"
    
    # Combinar la primera celda para "TAMAÑO DEL LOTE"
    cell_lote = table_lote.cell(0, 0)
    cell_lote.text = "TAMAÑO DEL LOTE"
    
    # Ajustar el ancho de las columnas
    cell_lote.width = Inches(4.5)
    table_lote.cell(0, 1).width = Inches(1.5)
    
    # Formatear celda de "TAMAÑO DEL LOTE"
    for paragraph in cell_lote.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Helvetica Neue"
            run.font.size = Pt(9)
    
    # Formatear celda del valor
    cell_valor = table_lote.cell(0, 1)
    cell_valor.text = "{{ TCantidad }}"
    for paragraph in cell_valor.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.name = "Helvetica Neue"
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ==================== OBSERVACIONES ====================
    p = doc.add_paragraph()
    p.add_run("OBSERVACIONES: ").bold = True
    p.add_run("La imagen amparada en el dictamen es una muestra de etiqueta que aplica para todos los modelos declarados en el presente dictamen lo anterior fue constatado durante la inspección.")
    for run in p.runs:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run("OBSERVACIONES: ").bold = True
    p.add_run("{{ obs }}")
    for run in p.runs:
        run.font.name = "Helvetica Neue"
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ==================== TABLA DE FIRMAS CON CUADROS ====================
    table_firmas = doc.add_table(rows=3, cols=2)
    table_firmas.style = "Table Grid"

    # Configurar ancho de columnas para firmas
    for row in table_firmas.rows:
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(3)

    # Configurar celdas de firma
    table_firmas.cell(0, 0).text = "{{ firma1 }}"
    table_firmas.cell(0, 1).text = "{{ firma2 }}"
    table_firmas.cell(1, 0).text = "{{ nfirma1 }}"
    table_firmas.cell(1, 1).text = "{{ nfirma2 }}"
    table_firmas.cell(2, 0).text = "Nombre del Inspector"
    table_firmas.cell(2, 1).text = "Nombre del responsable de\nsupervisión UI"

    # Aplicar formato a todas las celdas de la tabla de firmas
    for row_idx, row in enumerate(table_firmas.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = "Helvetica Neue"
                    run.font.size = Pt(9)
                    # Los títulos (tercera fila) en negrita
                    if row_idx == 2:
                        run.font.bold = True

    # ==================== GUARDAR ====================
    output_path = "Dictamen_plantilla.docx"
    try:
        doc.save(output_path)
        print(f"✅ Plantilla final creada correctamente: {output_path}")
        print("📝 La plantilla incluye los siguientes campos a reemplazar:")
        campos = [
            "{{ fverificacion }}", "{{ femision }}", "{{ cliente }}", "{{ rfc }}",
            "{{ producto }}", "{{ pedimento }}", "{{ fverificacionlarga }}", 
            "{{ capitulo }}", "{{ norma }}", "{{ normades }}",
            "{{ rowMarca }}", "{{ rowCodigo }}", "{{ rowFactura }}", "{{ rowCantidad }}",
            "{{ TCantidad }}", "{{ obs }}", "{{ firma1 }}", "{{ firma2 }}",
            "{{ nfirma1 }}", "{{ nfirma2 }}"
        ]
        for campo in campos:
            print(f"   - {campo}")
        print("\n✅ Características implementadas:")
        print("   - Cuadro de fechas en el encabezado")
        print("   - Cuadro de firmas en el pie de página") 
        print("   - Tabla de productos con bordes")
        print("   - Tamaño del lote con celda combinada")
        print("   - Imagen de fondo detrás del texto")
    except Exception as e:
        print(f"❌ Error al guardar la plantilla: {e}")

if __name__ == "__main__":
    crear_plantilla_con_fondo()