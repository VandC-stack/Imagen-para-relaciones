# -- AXO IMAGENES PARA LOS DICTAMENES DE AXO -- # 
import os
import json
import pandas as pd
from tkinter import filedialog, Tk
from docx import Document
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro
from main import obtener_rutas, insertar_imagen_con_transparencia, APPDATA_DIR

INDEX_FILE = os.path.join(APPDATA_DIR, "index_indice.json")
IMG_EXTS = [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif"]


def seleccionar_excel():
    Tk().withdraw()
    return filedialog.askopenfilename(
        title="Selecciona el archivo Excel para construir el índice",
        filetypes=[("Excel Files", "*.xlsx *.xlsm *.xls *.xlsb")]
    )


def construir_indice_desde_excel(ruta_excel):
    ext = os.path.splitext(ruta_excel)[1].lower()
    engine = "pyxlsb" if ext == ".xlsb" else None
    df = pd.read_excel(ruta_excel, sheet_name="CONCENTRADO", engine=engine)

    indice = {}

    for _, row in df.iterrows():
        try:
            codigo = str(row.iloc[0]).strip()
            destino = str(row.iloc[1]).strip()
        except Exception:
            continue

        if not codigo or not destino or codigo.lower() == "nan" or destino.lower() == "nan":
            continue

        if "código" in codigo.lower() or "sku" in codigo.lower():
            continue

        indice[codigo] = destino

    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(indice, f, ensure_ascii=False, indent=4)

    return indice


def extraer_codigos_tabla(doc):
    codigos = []
    for tabla in doc.tables:
        encabezados = [c.text.strip().upper() for c in tabla.rows[0].cells]
        if any("CODIGO" in h.replace("Ó", "O") or "SKU" in h or "CLAVE" in h for h in encabezados):
            idx = 0
            for i, h in enumerate(encabezados):
                h_norm = h.replace("Ó", "O")
                if "CODIGO" in h_norm or "SKU" in h_norm or "CLAVE" in h_norm:
                    idx = i
                    break
            for fila in tabla.rows[1:]:
                texto = fila.cells[idx].text.strip()
                if texto and any(c.isdigit() for c in texto) and len(texto) >= 6:
                    codigos.append(texto)
            break
    return codigos


def buscar_destino(ruta_base, destino):
    destino = destino.strip()
    
    base, ext = os.path.splitext(destino)

    if ext.lower() in IMG_EXTS:
        for archivo in os.listdir(ruta_base):
            if archivo.lower() == destino.lower():
                return "imagen", os.path.join(ruta_base, archivo)

    nombre_base = base if ext.lower() in IMG_EXTS else destino
    for archivo in os.listdir(ruta_base):
        archivo_base, archivo_ext = os.path.splitext(archivo)
        if archivo_base.lower() == nombre_base.lower() and archivo_ext.lower() in IMG_EXTS:
            return "imagen", os.path.join(ruta_base, archivo)

    carpeta_buscada = nombre_base
    for item in os.listdir(ruta_base):
        if os.path.isdir(os.path.join(ruta_base, item)) and item.lower() == carpeta_buscada.lower():
            return "carpeta", os.path.join(ruta_base, item)
    
    return None, None


def procesar_doc_con_indice(ruta_doc, ruta_imagenes, indice):
    doc = Document(ruta_doc)
    codigos = extraer_codigos_tabla(doc)

    fallo_registrado = False
    imagenes_insertadas = 0

    if not codigos:
        if not fallo_registrado:
            registrar_fallo(os.path.basename(ruta_doc))
            fallo_registrado = True
        return

    for p in doc.paragraphs:
        if "${etiqueta1}" in (p.text or ""):
            p.clear()
            run = p.add_run()

            for codigo in codigos:
                if codigo not in indice:
                    continue  # si no existe en el índice, probaremos con el siguiente

                destino = indice[codigo]
                tipo, ruta = buscar_destino(ruta_imagenes, destino)

                if tipo == "imagen":
                    insertar_imagen_con_transparencia(run, ruta)
                    imagenes_insertadas += 1

                elif tipo == "carpeta":
                    for archivo in os.listdir(ruta):
                        if os.path.splitext(archivo)[1].lower() in IMG_EXTS:
                            insertar_imagen_con_transparencia(run, os.path.join(ruta, archivo))
                            imagenes_insertadas += 1

                else:
                    if not fallo_registrado:
                        registrar_fallo(os.path.basename(ruta_doc))
                        fallo_registrado = True

            break

    if imagenes_insertadas == 0 and not fallo_registrado:
        registrar_fallo(os.path.basename(ruta_doc))
        fallo_registrado = True

    doc.save(ruta_doc)
    print(f"Documento actualizado: {ruta_doc}")

def procesar_indice():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    excel = seleccionar_excel()
    if not excel:
        raise Exception("No se seleccionó un archivo Excel para el modo Pegado por Índice.")

    print("Construyendo índice desde Excel...")
    indice = construir_indice_desde_excel(excel)
    print("Índice generado correctamente.")

    archivos = [f for f in os.listdir(ruta_docs) if f.endswith(".docx") and not f.startswith("~$")]

    for archivo in archivos:
        procesar_doc_con_indice(os.path.join(ruta_docs, archivo), ruta_imgs, indice)

    mostrar_registro()

    log = os.path.abspath("documentos_sin_imagenes.txt")
    if os.path.exists(log):
        os.startfile(log)

