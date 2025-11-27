import os
from docx import Document
from main import obtener_rutas, indexar_imagenes, buscar_imagen_index, insertar_imagen_con_transparencia
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE

def procesar_simple():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    index = indexar_imagenes(ruta_imgs)
    archivos = [f for f in os.listdir(ruta_docs) if f.endswith(".docx") and not f.startswith("~$")]

    for archivo in archivos:
        ruta_doc = os.path.join(ruta_docs, archivo)
        print(f"Procesando documento (modo simple): {ruta_doc}")
        doc = Document(ruta_doc)

        imagen_insertada = False
        usadas_paths = set()
        usadas_bases = set()

        from main import extraer_codigos, normalizar_cadena_alnum_mayus, norm_path_key
        codigos = extraer_codigos(doc)

        if not codigos:
            registrar_fallo(archivo)
            continue

        for p in doc.paragraphs:
            if "${etiqueta1}" in (p.text or ""):
                p.clear()
                run = p.add_run()

                for codigo in codigos:
                    img_path = buscar_imagen_index(index, codigo, usadas_paths, usadas_bases)
                    if img_path:
                        usadas_paths.add(norm_path_key(img_path))
                        usar_base = normalizar_cadena_alnum_mayus(os.path.splitext(os.path.basename(img_path))[0])
                        usadas_bases.add(usar_base)
                        insertar_imagen_con_transparencia(run, img_path)
                        imagen_insertada = True

                break

        if not imagen_insertada:
            registrar_fallo(archivo)

        doc.save(ruta_doc)
        print(f"Documento actualizado: {ruta_doc}")

    mostrar_registro()
    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)
