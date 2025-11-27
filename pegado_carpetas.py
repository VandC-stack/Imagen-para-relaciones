import os
from docx import Document
from main import (
    obtener_rutas,
    insertar_imagen_con_transparencia,
    extraer_codigos,
    normalizar_cadena_alnum_mayus,
)
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE

IMG_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


def construir_indice_carpetas(ruta_imgs):
    """
    Crea un índice de carpetas:
        clave normalizada (solo letras/números, mayúsculas) -> [rutas de carpeta]
    Esto permite que códigos como 'KI1545138' encuentren carpetas llamadas 'KI154-5138'.
    """
    indice = {}

    for nombre in os.listdir(ruta_imgs):
        ruta = os.path.join(ruta_imgs, nombre)
        if not os.path.isdir(ruta):
            continue

        clave = normalizar_cadena_alnum_mayus(nombre)
        if not clave:
            continue

        if clave not in indice:
            indice[clave] = []
        indice[clave].append(ruta)

    print(f"Índice de carpetas construido con {len(indice)} claves.")
    return indice


def procesar_carpetas():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    # Construimos un índice de carpetas una sola vez
    carpetas_index = construir_indice_carpetas(ruta_imgs)

    # Tomamos todos los .docx válidos
    archivos = [
        f for f in os.listdir(ruta_docs)
        if f.endswith(".docx") and not f.startswith("~$")
    ]

    for archivo in archivos:
        ruta_doc = os.path.join(ruta_docs, archivo)
        print(f"Procesando documento (modo carpetas): {ruta_doc}")
        doc = Document(ruta_doc)

        imagen_insertada = False
        codigos = extraer_codigos(doc)

        if not codigos:
            print("  No se encontraron códigos en el documento.")
            registrar_fallo(archivo)
            continue

        for p in doc.paragraphs:
            if "${etiqueta1}" in (p.text or ""):
                p.clear()
                run = p.add_run()

                for codigo in codigos:
                    # El código ya viene normalizado, pero por seguridad lo normalizamos otra vez
                    clave = normalizar_cadena_alnum_mayus(codigo)
                    if not clave:
                        continue

                    carpetas = carpetas_index.get(clave, [])
                    if not carpetas:
                        print(f"  No se encontró carpeta para código '{codigo}' (clave '{clave}').")
                        continue

                    for carpeta_codigo in carpetas:
                        for archivo_img in os.listdir(carpeta_codigo):
                            ext = os.path.splitext(archivo_img)[1].lower()
                            if ext in IMG_EXTS:
                                img_path = os.path.join(carpeta_codigo, archivo_img)
                                insertar_imagen_con_transparencia(run, img_path)
                                imagen_insertada = True
                                print(f"  Imagen insertada: {img_path}")

                break  # solo una vez por ${etiqueta1}

        if not imagen_insertada:
            registrar_fallo(archivo)

        doc.save(ruta_doc)
        print(f"Documento actualizado: {ruta_doc}")

    mostrar_registro()
    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)