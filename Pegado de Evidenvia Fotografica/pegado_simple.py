import os
from docx import Document
from main import (
    obtener_rutas,
    indexar_imagenes,
    buscar_imagen_index,
    insertar_imagen_con_transparencia,
    extraer_codigos,
    extraer_codigos_pdf,
    normalizar_cadena_alnum_mayus,
    norm_path_key,
    insertar_imagenes_en_pdf_placeholder,
)
from plantillaPDF import cargar_tabla_relacion
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE


def buscar_imagen_index_all(index, codigo_canonico, usadas_paths, usadas_bases):
    """Devuelve una lista de rutas de imagen en el índice que correspondan
    al código dado, incluyendo variantes como '1234(2)', '1234-2', '1234_2'.
    Respeta `usadas_paths` y `usadas_bases` para evitar duplicados.
    """
    code = normalizar_cadena_alnum_mayus(codigo_canonico)
    if not code:
        return []

    matches = []

    # Buscar coincidencias exactas y parciales
    for it in index:
        try:
            bn = it.get('base_norm') or ''
            bcore = it.get('base_core_norm') or ''
            path_key = norm_path_key(it.get('path') or '')
            if not bn or path_key in usadas_paths or bn in usadas_bases:
                continue
            # aceptar coincidencia si base_norm o base_core_norm empata
            if bn == code or bcore == code:
                matches.append(it['path'])
                continue
            # coincidencias parciales: incluir variantes donde bn contiene code
            if code in bn or bn.startswith(code) or bn.endswith(code) or code in bcore or bcore.startswith(code) or bcore.endswith(code):
                matches.append(it['path'])
        except Exception:
            continue

    # Deduplicado y mantener orden
    seen = set()
    out = []
    for p in matches:
        k = norm_path_key(p)
        if k in seen:
            continue
        seen.add(k)
        out.append(p)
    return out


def procesar_simple():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    # Índice normal de imágenes para el modo simple
    index = indexar_imagenes(ruta_imgs)

    # Ahora modo simple procesa Word y PDF
    archivos = [
        f for f in os.listdir(ruta_docs)
        if (f.lower().endswith(".docx") or f.lower().endswith(".pdf")) and not f.startswith("~$")
    ]

    for archivo in archivos:
        ruta_doc = os.path.join(ruta_docs, archivo)
        ext = os.path.splitext(archivo)[1].lower()

        # ========================================
        # WORD
        # ========================================
        if ext == ".docx":
            print(f"Procesando documento (modo simple DOCX): {ruta_doc}")

            doc = Document(ruta_doc)
            codigos = extraer_codigos(doc)

            # Filtrar códigos contra la tabla de relación (si existe)
            try:
                df_rel = cargar_tabla_relacion()
                valid_codes = set()
                for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
                    if col in df_rel.columns:
                        for v in df_rel[col].astype(str).fillna(""):
                            valid_codes.add(normalizar_cadena_alnum_mayus(v))
                        break
            except Exception:
                valid_codes = None

            if codigos:
                if valid_codes is not None:
                    codigos = [c for c in codigos if normalizar_cadena_alnum_mayus(c) in valid_codes]

            if not codigos:
                registrar_fallo(archivo)
                continue

            imagen_insertada = False
            usadas_paths = set()
            usadas_bases = set()

            for p in doc.paragraphs:
                 if "${imagen}" in (p.text or ""):
                    p.clear()
                    run = p.add_run()

                    for codigo in codigos:
                        img_paths = buscar_imagen_index_all(index, codigo, usadas_paths, usadas_bases)
                        if img_paths:
                            for img_path in img_paths:
                                kp = norm_path_key(img_path)
                                if kp in usadas_paths:
                                    continue
                                usadas_paths.add(kp)
                                usar_base = normalizar_cadena_alnum_mayus(os.path.splitext(os.path.basename(img_path))[0])
                                usadas_bases.add(usar_base)
                                insertar_imagen_con_transparencia(run, img_path)
                                imagen_insertada = True

                    break

            if not imagen_insertada:
                registrar_fallo(archivo)

            doc.save(ruta_doc)
            print(f"Documento DOCX actualizado: {ruta_doc}")
            continue

        # ========================================
        # PDF
        # ========================================
        if ext == ".pdf":
            print(f"Procesando documento (modo simple PDF): {ruta_doc}")

            codigos = extraer_codigos_pdf(ruta_doc)

            # Filtrar códigos contra la tabla de relación (si existe)
            try:
                df_rel = cargar_tabla_relacion()
                valid_codes = set()
                for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
                    if col in df_rel.columns:
                        for v in df_rel[col].astype(str).fillna(""):
                            valid_codes.add(normalizar_cadena_alnum_mayus(v))
                        break
            except Exception:
                valid_codes = None

            if codigos:
                if valid_codes is not None:
                    codigos = [c for c in codigos if normalizar_cadena_alnum_mayus(c) in valid_codes]

            if not codigos:
                registrar_fallo(archivo)
                continue

            rutas_imagenes = []
            usadas_paths = set()
            usadas_bases = set()

            for codigo in codigos:
                img_paths = buscar_imagen_index_all(index, codigo, usadas_paths, usadas_bases)
                if img_paths:
                    for img_path in img_paths:
                        kp = norm_path_key(img_path)
                        if kp in usadas_paths:
                            continue
                        rutas_imagenes.append(img_path)
                        usadas_paths.add(kp)
                        usar_base = normalizar_cadena_alnum_mayus(os.path.splitext(os.path.basename(img_path))[0])
                        usadas_bases.add(usar_base)

            if not rutas_imagenes:
                registrar_fallo(archivo)
                continue

            exito = insertar_imagenes_en_pdf_placeholder(ruta_doc, rutas_imagenes)
            if not exito:
                registrar_fallo(archivo)

    mostrar_registro()
    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)  