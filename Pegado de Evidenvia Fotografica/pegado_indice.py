import os
import json
import pandas as pd
import re
from datetime import datetime
from tkinter import filedialog, Tk
from docx import Document
from registro_fallos import registrar_fallo, limpiar_registro, mostrar_registro, LOG_FILE
from main import (
    obtener_rutas,
    insertar_imagen_con_transparencia,
    APPDATA_DIR,
    extraer_codigos_pdf,
    insertar_imagenes_en_pdf_placeholder,
)
from main import normalizar_cadena_alnum_mayus
from plantillaPDF import cargar_tabla_relacion

INDEX_FILE = os.path.join(APPDATA_DIR, "index_indice.json")
IMG_EXTS = [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif"]

# Debug logger for pegado indice
DEBUG_INDEX_LOG = os.path.join(APPDATA_DIR, "pegado_indice_debug.log")
def _log_index(msg: str):
    try:
        with open(DEBUG_INDEX_LOG, 'a', encoding='utf-8') as lf:
            lf.write(f"[{datetime.now().isoformat()}] {msg}\n")
    except Exception:
        try:
            print("[pegado_indice]", msg)
        except Exception:
            pass


def seleccionar_excel():
    Tk().withdraw()
    return filedialog.askopenfilename(
        title="Selecciona el archivo Excel para construir el índice",
        filetypes=[("Excel Files", "*.xlsx *.xlsm *.xls *.xlsb")]
    )


def construir_indice_desde_excel(ruta_excel):
    ext = os.path.splitext(ruta_excel)[1].lower()
    engine = "pyxlsb" if ext == ".xlsb" else None
    try:
        df = pd.read_excel(ruta_excel, sheet_name="CONCENTRADO", engine=engine)
    except Exception as e:
        msg = str(e).lower()
        if ext == ".xlsb" and ("pyxlsb" in msg or "missing optional dependency" in msg):
            raise Exception("Missing optional dependency 'pyxlsb'.") from e
        raise

    indice = {}

    # Cargar tabla de relación para filtrar códigos válidos
    try:
        df_rel = cargar_tabla_relacion()
        valid_codes = set()
        for col in ("CODIGO","CODIGOS","CODE","SKU","CLAVE"):
            if col in df_rel.columns:
                for v in df_rel[col].astype(str).fillna(""):
                    valid_codes.add(normalizar_cadena_alnum_mayus(v))
                break
    except Exception:
        valid_codes = None

    for _, row in df.iterrows():
        try:
            codigo = str(row.iloc[0]).strip()
            destino = str(row.iloc[1]).strip()
        except Exception:
            continue

        if not codigo or not destino or codigo.lower() == "nan" or destino.lower() == "nan":
            continue

        if "código" in codigo.lower() or "sku" in codigo.lower():
            continue

        canon = normalizar_cadena_alnum_mayus(codigo)
        # Si disponemos de la tabla de relación, solo añadimos códigos que estén en ella
        if valid_codes is not None and canon not in valid_codes:
            continue

        indice[canon] = destino

    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(indice, f, ensure_ascii=False, indent=4)
    _log_index(f"Indice construido: {len(indice)} entries; index_file={INDEX_FILE}")

    return indice


def extraer_codigos_tabla(doc):
    codigos = []
    for tabla in doc.tables:
        if not tabla.rows:
            continue

        encabezados = [c.text.strip().upper() for c in tabla.rows[0].cells]
        if any("CODIGO" in h.replace("Ó", "O") or "SKU" in h or "CLAVE" in h for h in encabezados):
            idx = 0
            for i, h in enumerate(encabezados):
                h_norm = h.replace("Ó", "O")
                if "CODIGO" in h_norm or "SKU" in h_norm or "CLAVE" in h_norm:
                    idx = i
                    break

            for fila in tabla.rows[1:]:
                texto = (fila.cells[idx].text or "").strip()
                if not texto:
                    continue

                canon = "".join(ch for ch in texto if ch.isalnum())
                if not canon:
                    continue

                tiene_letras = any(c.isalpha() for c in canon)
                tiene_digitos = any(c.isdigit() for c in canon)

                if not (tiene_letras and tiene_digitos):
                    continue

                codigos.append(texto)

            break
    return codigos


def buscar_destino(ruta_base, destino):
    destino = destino.strip()

    base, ext = os.path.splitext(destino)

    if ext.lower() in IMG_EXTS:
        for archivo in os.listdir(ruta_base):
            if archivo.lower() == destino.lower():
                return "imagen", os.path.join(ruta_base, archivo)

    nombre_base = base if ext.lower() in IMG_EXTS else destino
    for archivo in os.listdir(ruta_base):
        archivo_base, archivo_ext = os.path.splitext(archivo)
        if archivo_base.lower() == nombre_base.lower() and archivo_ext.lower() in IMG_EXTS:
            return "imagen", os.path.join(ruta_base, archivo)

    carpeta_buscada = nombre_base
    for item in os.listdir(ruta_base):
        if os.path.isdir(os.path.join(ruta_base, item)) and item.lower() == carpeta_buscada.lower():
            return "carpeta", os.path.join(ruta_base, item)

    return None, None


def procesar_doc_con_indice_docx(ruta_doc, ruta_imagenes, indice):
    doc = Document(ruta_doc)
    codigos = extraer_codigos_tabla(doc)

    fallo_registrado = False
    imagenes_insertadas = 0

    if not codigos:
        if not fallo_registrado:
            registrar_fallo(os.path.basename(ruta_doc))
            fallo_registrado = True
        return

    _log_index(f"Procesando DOCX: {ruta_doc}; codigos_found={len(codigos)}; ruta_imagenes={ruta_imagenes}")
    for p in doc.paragraphs:
        txt = (p.text or "")
        # Accept case-variants like ${IMAGEN} or ${imagen} (allow spaces inside braces)
        if re.search(r"\$\{\s*imagen\s*\}", txt, flags=re.IGNORECASE):
            p.clear()
            run = p.add_run()

            for codigo in codigos:
                canon = normalizar_cadena_alnum_mayus(codigo)
                if canon not in indice:
                    continue

                destino = indice[canon]
                _log_index(f"Codigo {canon} -> destino {destino}")
                tipo, ruta = buscar_destino(ruta_imagenes, destino)
                _log_index(f"buscar_destino -> tipo={tipo} ruta={ruta}")

                if tipo == "imagen":
                    insertar_imagen_con_transparencia(run, ruta)
                    imagenes_insertadas += 1

                elif tipo == "carpeta":
                    for archivo in os.listdir(ruta):
                        if os.path.splitext(archivo)[1].lower() in IMG_EXTS:
                            insertar_imagen_con_transparencia(run, os.path.join(ruta, archivo))
                            imagenes_insertadas += 1

                else:
                    if not fallo_registrado:
                        registrar_fallo(os.path.basename(ruta_doc))
                        fallo_registrado = True

            break

    if imagenes_insertadas == 0 and not fallo_registrado:
        registrar_fallo(os.path.basename(ruta_doc))
        fallo_registrado = True

    doc.save(ruta_doc)
    print(f"Documento actualizado: {ruta_doc}")


def procesar_doc_con_indice_pdf(ruta_doc, ruta_imagenes, indice):
    codigos = extraer_codigos_pdf(ruta_doc)

    fallo_registrado = False
    imagenes_insertadas = 0

    if not codigos:
        if not fallo_registrado:
            registrar_fallo(os.path.basename(ruta_doc))
            fallo_registrado = True
        return

    rutas_imagenes = []

    _log_index(f"Procesando PDF: {ruta_doc}; codigos_found={len(codigos)}; ruta_imagenes={ruta_imagenes}")
    for codigo in codigos:
        canon = normalizar_cadena_alnum_mayus(codigo)
        if canon not in indice:
            continue

        destino = indice[canon]
        _log_index(f"Codigo {canon} -> destino {destino}")
        tipo, ruta = buscar_destino(ruta_imagenes, destino)
        _log_index(f"buscar_destino -> tipo={tipo} ruta={ruta}")

        if tipo == "imagen":
            rutas_imagenes.append(ruta)
            imagenes_insertadas += 1

        elif tipo == "carpeta":
            for archivo in os.listdir(ruta):
                if os.path.splitext(archivo)[1].lower() in IMG_EXTS:
                    rutas_imagenes.append(os.path.join(ruta, archivo))
                    imagenes_insertadas += 1

        else:
            if not fallo_registrado:
                registrar_fallo(os.path.basename(ruta_doc))
                fallo_registrado = True

    if imagenes_insertadas == 0 and not fallo_registrado:
        registrar_fallo(os.path.basename(ruta_doc))
        fallo_registrado = True
        return

    exito = insertar_imagenes_en_pdf_placeholder(ruta_doc, rutas_imagenes)
    if not exito and not fallo_registrado:
        registrar_fallo(os.path.basename(ruta_doc))


def procesar_doc_con_indice(ruta_doc, ruta_imagenes, indice):
    ext = os.path.splitext(ruta_doc)[1].lower()
    if ext == ".docx":
        procesar_doc_con_indice_docx(ruta_doc, ruta_imagenes, indice)
    elif ext == ".pdf":
        procesar_doc_con_indice_pdf(ruta_doc, ruta_imagenes, indice)


def procesar_indice():
    limpiar_registro()

    ruta_docs, ruta_imgs = obtener_rutas()
    if not ruta_docs or not ruta_imgs:
        return

    excel = seleccionar_excel()
    if not excel:
        raise Exception("No se seleccionó un archivo Excel para el modo Pegado por Índice.")

    print("Construyendo índice desde Excel...")
    indice = construir_indice_desde_excel(excel)
    print("Índice generado correctamente.")

    archivos = [
        f for f in os.listdir(ruta_docs)
        if (f.endswith(".docx") or f.endswith(".pdf")) and not f.startswith("~$")
    ]

    for archivo in archivos:
        procesar_doc_con_indice(os.path.join(ruta_docs, archivo), ruta_imgs, indice)

    mostrar_registro()

    if os.path.exists(LOG_FILE):
        os.startfile(LOG_FILE)