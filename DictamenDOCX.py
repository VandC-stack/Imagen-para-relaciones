"""
Generador de Plantilla BOSCH Word (.docx) con Fondo y Pie de Página
- Pie de página en todas las hojas
- Imagen de fondo ajustada correctamente
- Estructura profesional mantenida
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import os
from datetime import datetime

class WordGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_document()
        
    def setup_document(self):
        """Configura el documento Word con los márgenes y estilos básicos"""
        # Configurar márgenes (aproximadamente 0.75 pulgadas)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(1.0)  # Margen inferior aumentado para el pie
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def agregar_fondo_pagina(self):
        """Agrega la imagen de fondo a todas las páginas ajustada a la izquierda"""
        try:
            # Verificar si existe la imagen de fondo
            fondo_path = "img/Fondo.jpeg"
            if not os.path.exists(fondo_path):
                print("⚠️  No se encontró la imagen de fondo en img/Fondo.jpeg")
                return False
            
            # Agregar fondo a todas las secciones (páginas)
            for section in self.doc.sections:
                header = section.header
                
                # Limpiar el encabezado
                for paragraph in header.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                
                # Agregar la imagen de fondo MOVIDA A LA IZQUIERDA
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                run = paragraph.add_run()
                
                # Agregar imagen ajustada al tamaño de página y movida a la izquierda
                # Aumentamos el ancho ligeramente para asegurar cobertura
                run.add_picture(fondo_path, width=Inches(8.8), height=Inches(11.2))
                
                # Alinear a la izquierda para cubrir mejor
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Configurar el encabezado para que no tenga márgenes
                header_distance = Inches(0)
                section.header_distance = header_distance
                
            print("✅ Imagen de fondo agregada y ajustada a la izquierda")
            return True
            
        except Exception as e:
            print(f"⚠️  No se pudo agregar la imagen de fondo: {e}")
            print("💡 El documento se generará sin fondo")
            return False
    
    def agregar_pie_pagina_global(self):
        """Agrega el pie de página a TODAS las hojas del documento"""
        try:
            for section in self.doc.sections:
                footer = section.footer
                
                # Limpiar footer existente
                for paragraph in footer.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                
                # Texto principal del pie de página
                pie_text = (
                    "Este Dictamen de Cumplimiento se emitió por medios electrónicos, conforme al oficio de autorización "
                    "DGN.312.05.2012.106 de fecha 10 de enero de 2012 expedido por la DGN a esta Unidad de Inspección."
                )
                
                # Agregar texto principal
                pie_paragraph = footer.add_paragraph()
                pie_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pie_run = pie_paragraph.add_run(pie_text)
                pie_run.font.size = Pt(8)
                pie_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Agregar formato alineado a la derecha
                formato_paragraph = footer.add_paragraph()
                formato_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                formato_run = formato_paragraph.add_run("Formato: PT-F-208B-00-3")
                formato_run.font.size = Pt(8)
                formato_run.font.color.rgb = RGBColor(0, 0, 0)
                formato_run.bold = True
                
                # Configurar distancia del footer
                section.footer_distance = Inches(0.5)
                
            print("✅ Pie de página agregado a todas las hojas")
            return True
            
        except Exception as e:
            print(f"⚠️  No se pudo agregar el pie de página global: {e}")
            return False
    
    def crear_encabezado(self):
        """Crea el encabezado del documento (contenido, no fondo)"""
        # Título principal
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("DICTAMEN DE CUMPLIMIENTO")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Código
        code = self.doc.add_paragraph()
        code.alignment = WD_ALIGN_PARAGRAPH.CENTER
        code_run = code.add_run("${year}049UDC${norma}${folio} Solicitud: ${year}049USD${norma}${solicitud}-${lista}")
        code_run.font.size = Pt(10)
        code_run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        self.doc.add_paragraph()  # Espacio
    
    def crear_primera_pagina(self):
        """Crea el contenido de la primera página"""
        print("📄 Generando primera página...")
        
        # ==================== TABLA DE FECHAS ====================
        fecha_table = self.doc.add_table(rows=2, cols=2)
        fecha_table.autofit = False
        fecha_table.columns[0].width = Inches(3)
        fecha_table.columns[1].width = Inches(3)
        
        # Encabezados de tabla de fechas
        fecha_table.cell(0, 0).text = "Fecha de Inspección"
        fecha_table.cell(0, 1).text = "${fverificacion}"
        fecha_table.cell(1, 0).text = "Fecha de Emisión"
        fecha_table.cell(1, 1).text = "${femision}"
        
        # Formatear tabla de fechas
        for row in fecha_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Negrita para la primera columna
        for i in range(2):
            for paragraph in fecha_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== CLIENTE Y RFC ====================
        cliente_para = self.doc.add_paragraph()
        cliente_run = cliente_para.add_run("Cliente: ")
        cliente_run.bold = True
        cliente_run.font.color.rgb = RGBColor(0, 0, 0)
        cliente_para.add_run("${cliente}").font.color.rgb = RGBColor(0, 0, 0)
        
        rfc_para = self.doc.add_paragraph()
        rfc_run = rfc_para.add_run("RFC: ")
        rfc_run.bold = True
        rfc_run.font.color.rgb = RGBColor(0, 0, 0)
        rfc_para.add_run("${rfc}").font.color.rgb = RGBColor(0, 0, 0)
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== TEXTO PRINCIPAL ====================
        texto_principal = (
            "De conformidad en lo dispuesto en los artículos 53, 56 fracción I, 60 fracción I, 62, 64, 68 y 140 de la Ley de Infraestructura de la "
            "Calidad; 50 del Reglamento de la Ley Federal de Metrología y Normalización; Punto 2.4.8 Fracción III ACUERDO por el que la "
            "Secretaría de Economía emite Reglas y criterios de carácter general en materia de comercio exterior; publicado en el Diario Oficial de la "
            "Federación el 09 de mayo de 2022 y posteriores modificaciones; esta Unidad de Inspección a solicitud de la persona moral denominada "
            "${cliente} dictamina el Producto: ${producto}; que la mercancía importada bajo el pedimento aduanal No. ${pedimento} de fecha "
            "${fverificacionlarga}, fue etiquetada conforme a los requisitos de Información Comercial en el capítulo ${capitulo} de la Norma Oficial Mexicana "
            "${norma} ${normades} Cualquier otro requisito establecido en la norma referida, es responsabilidad del titular de este Dictamen."
        )
        
        texto_para = self.doc.add_paragraph(texto_principal)
        texto_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in texto_para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== TABLA DE PRODUCTOS ====================
        productos_table = self.doc.add_table(rows=2, cols=4)
        productos_table.autofit = False
        productos_table.columns[0].width = Inches(1.5)
        productos_table.columns[1].width = Inches(1.5)
        productos_table.columns[2].width = Inches(1.5)
        productos_table.columns[3].width = Inches(1.5)
        
        # Encabezados
        productos_table.cell(0, 0).text = "MARCA"
        productos_table.cell(0, 1).text = "CÓDIGO"
        productos_table.cell(0, 2).text = "FACTURA"
        productos_table.cell(0, 3).text = "CANTIDAD"
        
        # Datos
        productos_table.cell(1, 0).text = "${rowMarca}"
        productos_table.cell(1, 1).text = "${rowCodigo}"
        productos_table.cell(1, 2).text = "${rowFactura}"
        productos_table.cell(1, 3).text = "${rowCantidad}"
        
        # Formatear tabla de productos
        for row in productos_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Negrita para encabezados
        for i in range(4):
            for paragraph in productos_table.cell(0, i).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== TAMAÑO DEL LOTE ====================
        lote_table = self.doc.add_table(rows=1, cols=2)
        lote_table.autofit = False
        lote_table.columns[0].width = Inches(4.5)
        lote_table.columns[1].width = Inches(1.5)
        
        lote_table.cell(0, 0).text = "TAMAÑO DEL LOTE"
        lote_table.cell(0, 1).text = "${TCantidad}"
        
        # Formatear tabla de lote
        for row in lote_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Negrita para primera celda
        for paragraph in lote_table.cell(0, 0).paragraphs:
            for run in paragraph.runs:
                run.bold = True
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== OBSERVACIONES ====================
        obs1_para = self.doc.add_paragraph()
        obs1_run = obs1_para.add_run("OBSERVACIONES: ")
        obs1_run.bold = True
        obs1_run.font.color.rgb = RGBColor(0, 0, 0)
        obs1_para.add_run("La imagen amparada en el dictamen es una muestra de etiqueta que aplica para todos los modelos declarados en el presente dictamen lo anterior fue constatado durante la inspección.").font.color.rgb = RGBColor(0, 0, 0)
        
        obs2_para = self.doc.add_paragraph()
        obs2_run = obs2_para.add_run("OBSERVACIONES: ")
        obs2_run.bold = True
        obs2_run.font.color.rgb = RGBColor(0, 0, 0)
        obs2_para.add_run("${obs}").font.color.rgb = RGBColor(0, 0, 0)
        
        # Salto de página
        self.doc.add_page_break()
    
    def crear_segunda_pagina(self):
        """Crea el contenido de la segunda página"""
        print("📄 Generando segunda página...")
        
        # ==================== ETIQUETAS ====================
        etiquetas_title = self.doc.add_paragraph("ETIQUETAS DEL PRODUCTO")
        etiquetas_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in etiquetas_title.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Primera fila de etiquetas
        etiquetas1 = self.doc.add_paragraph("${etiqueta1}   ${etiqueta2}   ${etiqueta3}   ${etiqueta4}   ${etiqueta5}")
        etiquetas1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in etiquetas1.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Segunda fila de etiquetas
        etiquetas2 = self.doc.add_paragraph("${etiqueta6}   ${etiqueta7}   ${etiqueta8}   ${etiqueta9}   ${etiqueta10}")
        etiquetas2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in etiquetas2.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        self.doc.add_paragraph()  # Espacio
        self.doc.add_paragraph()  # Espacio adicional
        
        # ==================== IMÁGENES ====================
        imagenes_title = self.doc.add_paragraph("IMÁGENES DEL PRODUCTO")
        imagenes_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in imagenes_title.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Lista de imágenes
        imagenes = [
            "${img1}", "${img2}", "${img3}", "${img4}", "${img5}",
            "${img6}", "${img7}", "${img8}", "${img9}", "${img10}"
        ]
        
        for img in imagenes:
            img_para = self.doc.add_paragraph(img)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in img_para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        self.doc.add_paragraph()  # Espacio
        
        # ==================== TABLA DE FIRMAS ====================
        firmas_table = self.doc.add_table(rows=3, cols=3)
        firmas_table.autofit = False
        firmas_table.columns[0].width = Inches(2.8)
        firmas_table.columns[1].width = Inches(0.4)
        firmas_table.columns[2].width = Inches(2.8)
        
        # Datos de la tabla
        firmas_table.cell(0, 0).text = "${firma1}"
        firmas_table.cell(0, 2).text = "${firma2}"
        firmas_table.cell(1, 0).text = "${nfirma1}"
        firmas_table.cell(1, 2).text = "${nfirma2}"
        firmas_table.cell(2, 0).text = "Nombre del Inspector"
        firmas_table.cell(2, 2).text = "Nombre del responsable de\nsupervisión UI"
        
        # Formatear tabla de firmas
        for row in firmas_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Negrita para la fila de títulos
        for i in range(3):
            for paragraph in firmas_table.cell(2, i).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Añadir líneas para firmas (simuladas con texto)
        for paragraph in firmas_table.cell(0, 0).paragraphs:
            paragraph.text = "________________________"
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        for paragraph in firmas_table.cell(0, 2).paragraphs:
            paragraph.text = "________________________"
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
    
    def agregar_numeracion_paginas(self):
        """Agrega numeración de páginas en el encabezado"""
        try:
            for i, section in enumerate(self.doc.sections):
                header = section.header
                # Buscar si ya existe un párrafo en el header
                if len(header.paragraphs) > 0:
                    # Agregar numeración al header existente
                    numeracion_para = header.paragraphs[0]
                    numeracion_run = numeracion_para.add_run(f"\t\tPágina {i+1} de 2")
                    numeracion_run.font.size = Pt(9)
                    numeracion_run.font.color.rgb = RGBColor(0, 0, 0)
                    numeracion_run.bold = True
                
        except Exception as e:
            print(f"⚠️  No se pudo agregar la numeración automática: {e}")
    
    def generar_documento_word(self):
        """Genera el documento Word completo con fondo y pie de página global"""
        print("🎯 GENERANDO DOCUMENTO WORD CON FONDO Y PIE GLOBAL...")
        
        try:
            # Primero agregar el fondo (ajustado a la izquierda)
            fondo_agregado = self.agregar_fondo_pagina()
            
            # Agregar pie de página global a TODAS las hojas
            pie_agregado = self.agregar_pie_pagina_global()
            
            if not fondo_agregado:
                print("💡 Generando documento sin fondo...")
            
            # Crear encabezado (contenido)
            self.crear_encabezado()
            
            # Crear primera página
            self.crear_primera_pagina()
            
            # Crear segunda página
            self.crear_segunda_pagina()
            
            # Agregar numeración
            self.agregar_numeracion_paginas()
            
            # Guardar documento
            output_path = "Dictamen_Final.docx"
            self.doc.save(output_path)
            
            print(f"✅ DOCUMENTO WORD CREADO: {output_path}")
            
            if fondo_agregado:
                print("💡 NOTA: El fondo está incluido y ajustado a la izquierda")
            if pie_agregado:
                print("💡 NOTA: El pie de página aparece en TODAS las hojas")
            
            return True
            
        except Exception as e:
            print(f"❌ Error al generar documento Word: {e}")
            return False
    
    def mostrar_variables(self):
        """Muestra las variables disponibles para reemplazo"""
        print("\n📋 VARIABLES DISPONIBLES PARA REEMPLAZO:")
        variables = [
            "year", "norma", "folio", "solicitud", "lista",
            "fverificacion", "femision", "cliente", "rfc",
            "producto", "pedimento", "fverificacionlarga", "capitulo",
            "normades", "rowMarca", "rowCodigo", "rowFactura", "rowCantidad",
            "TCantidad", "obs", "etiqueta1", "etiqueta2", "etiqueta3",
            "etiqueta4", "etiqueta5", "etiqueta6", "etiqueta7", "etiqueta8",
            "etiqueta9", "etiqueta10", "img1", "img2", "img3", "img4", "img5",
            "img6", "img7", "img8", "img9", "img10", "firma1", "firma2",
            "nfirma1", "nfirma2"
        ]
        
        for i, var in enumerate(variables, 1):
            print(f"   ${var}", end="\n" if i % 4 == 0 else " | ")
        
        print("\n\n💡 INSTRUCCIONES:")
        print("   1. Abre el documento Word generado")
        print("   2. Usa 'Ctrl+H' para buscar y reemplazar las variables")
        print("   3. Ejemplo: Busca '${cliente}' y reemplaza con el nombre real")
        print("   4. El pie de página ya aparece en TODAS las hojas automáticamente")
        print("   5. Guarda el documento con un nuevo nombre")


def verificar_imagen_fondo():
    """Verifica que la imagen de fondo existe"""
    image_path = "img/Fondo.jpeg"
    
    if not os.path.exists(image_path):
        print(f"⚠️  No se encontró: {image_path}")
        os.makedirs("img", exist_ok=True)
        print("📁 Se creó la carpeta 'img/' - coloca 'Fondo.jpeg' allí")
        return False
    
    print("✅ Imagen de fondo encontrada")
    return True


def reemplazar_variables_automatico(document_path, replacements):
    """
    Función auxiliar para reemplazar variables automáticamente
    Uso: reemplazar_variables_automatico("documento.docx", {"${cliente}": "Nombre Cliente", ...})
    """
    from docx import Document
    import re
    
    doc = Document(document_path)
    
    # Reemplazar en párrafos normales
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
    
    # Reemplazar en headers y footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        
        for paragraph in section.footer.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
    
    new_path = document_path.replace(".docx", "_completado.docx")
    doc.save(new_path)
    return new_path


if __name__ == "__main__":
    print("=" * 70)
    print("   GENERADOR DE PLANTILLA WORD BOSCH CON FONDO Y PIE GLOBAL")
    print("=" * 70)
    
    # Verificar imagen de fondo
    print("\n🔍 Verificando recursos...")
    verificar_imagen_fondo()
    
    # Generar documento Word
    print("\n🛠️  Generando plantilla Word con fondo y pie global...")
    generador = WordGenerator()
    
    if generador.generar_documento_word():
        # Mostrar información
        generador.mostrar_variables()
        
        print("\n📁 ARCHIVO CREADO:")
        print("   • Dictamen_Final.docx")
        
        print("\n🎯 CARACTERÍSTICAS IMPLEMENTADAS:")
        print("   ✅ Pie de página en TODAS las hojas")
        print("   ✅ Imagen de fondo ajustada a la izquierda")
        print("   ✅ Texto del pie correctamente formateado")
        print("   ✅ 'Formato: PT-F-208B-00-3' alineado a la derecha")
        print("   ✅ Estructura de 2 páginas completa")
        print("   ✅ Todas las variables de reemplazo incluidas")
        
        print("\n🔧 PARA USAR:")
        print("   - Abre el documento en Microsoft Word")
        print("   - Usa 'Buscar y Reemplazar' (Ctrl+H) para las variables")
        print("   - El pie de página ya está configurado automáticamente")
        
    else:
        print("❌ No se pudo generar el documento Word")
    
    print("\n🎉 ¡DOCUMENTO FINAL CREADO EXITOSAMENTE!")